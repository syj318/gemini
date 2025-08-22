# import os
# import streamlit as st
# import google.generativeai as genai
# from datetime import datetime
# from typing import Dict, List, Tuple

# from config import Config
# from utils import DB, UI
# from bexco_data import BEXCO_DATA, BEXCO_FAQ, get_bexco_info, search_bexco_info
# from add_partition import add_monthly_partition 


# # ======================= 벡스코 챗봇 =======================
# class BexcoChatbot:
#     @staticmethod
#     @st.cache_resource
#     def load_model():
#         try:
#             with open(Config.API_KEY_PATH, "r", encoding="utf-8") as f:
#                 api_key = f.read().strip()
#             genai.configure(api_key=api_key)
#             return genai.GenerativeModel(Config.MODEL_NAME)
#         except Exception as e:
#             st.error(f"모델 로드 실패: {e}")
#             st.stop()

#     def __init__(self):
#         self.model = BexcoChatbot.load_model()
#         self._ensure_session()

#     @staticmethod
#     def _ensure_session():
#         DB.init()
#         ss = st.session_state
#         ss.setdefault("session_id", os.urandom(8).hex())
#         ss.setdefault("chat_session", BexcoChatbot.load_model().start_chat(history=[]))
#         ss.setdefault("messages", [])
#         ss.setdefault("show_db_recent", False)

#     def _create_bexco_context(self, user_prompt: str) -> str:
#         """벡스코 정보를 포함한 컨텍스트 생성"""
#         context = f"사용자 질문: {user_prompt}\n\n간단하고 직접적으로 답변해주세요."
        
#         # 업로드된 파일이 있는 경우 파일 내용을 컨텍스트에 추가
#         if hasattr(st.session_state, 'uploaded_file_content') and st.session_state.uploaded_file_content:
#             context += f"""

# 파일 내용:
# {st.session_state.uploaded_file_content[:2000]}{'...' if len(st.session_state.uploaded_file_content) > 2000 else ''}
# """
        
#         return context

#     def reply(self, user_prompt: str) -> str:
#         try:
#             # 1. 먼저 데이터셋에서 정확한 답변을 찾아보기
#             search_results = search_bexco_info(user_prompt)
            
#             # 2. 데이터셋에 정확한 답변이 있는지 확인
#             if search_results:
#                 # 데이터셋에서 찾은 정보로 답변 구성
#                 context = f"""사용자 질문: {user_prompt}

# 관련 정보:
# {chr(10).join(search_results[:3])}

# 위 정보를 바탕으로 간단하게 답변해주세요."""
                
#                 # Gemini 모델로 데이터셋 기반 답변 생성
#                 stream = st.session_state.chat_session.send_message(context, stream=True)
#                 text = ""
#                 for chunk in stream:
#                     if hasattr(chunk, "text"):
#                         text += chunk.text
#                 return text
            
#             else:
#                 # 3. 데이터셋에 없는 질문인 경우 Gemini AI로 자유 답변 생성
#                 context = self._create_bexco_context(user_prompt)
                
#                 # Gemini 모델로 자유 답변 생성
#                 stream = st.session_state.chat_session.send_message(context, stream=True)
#                 text = ""
#                 for chunk in stream:
#                     if hasattr(chunk, "text"):
#                         text += chunk.text
#                 return text
                
#         except Exception as e:
#             return f"응답 생성 중 오류가 발생했습니다: {e}"


# # ======================= 메인 =======================
# def main():
#     st.set_page_config(
#         page_title="부산 벡스코 챗봇", 
#         layout="wide", 
#         initial_sidebar_state="expanded"
#     )
    
#     # ✅ 실행 시 매달 파티션 자동 생성
#     try:
#         add_monthly_partition()
#     except Exception as e:
#         st.warning(f"⚠️ 파티션 생성 중 오류 발생: {e}")
    
#     # 헤더
#     st.title("🏢 부산벡스코(BEXCO) 챗봇")
    
#     # CSS 스타일 적용
#     UI.css()
    
#     # 챗봇 초기화
#     bot = BexcoChatbot()

#     # ---------- 사이드바 정보 ----------
#     with st.sidebar:
#         st.header("📎 파일 업로드")
#         st.markdown("파일을 업로드하고 내용에 대해 질문하세요!")
        
#         # 파일 업로드
#         uploaded_file = st.file_uploader(
#             "파일 선택",
#             type=["pdf", "txt", "csv", "docx", "xlsx"],
#             help="PDF, TXT, CSV, DOCX, XLSX 파일을 업로드할 수 있습니다."
#         )
        
#         # 업로드된 파일 처리
#         if uploaded_file is not None:
#             file_content = ""
#             file_name = uploaded_file.name
            
#             try:
#                 if file_name.lower().endswith('.pdf'):
#                     import PyPDF2
#                     pdf_reader = PyPDF2.PdfReader(uploaded_file)
#                     for page in pdf_reader.pages:
#                         file_content += page.extract_text() + "\n"
                        
#                 elif file_name.lower().endswith('.txt'):
#                     file_content = uploaded_file.read().decode('utf-8')
                    
#                 elif file_name.lower().endswith('.csv'):
#                     import pandas as pd
#                     df = pd.read_csv(uploaded_file)
#                     file_content = f"CSV 파일 내용:\n{df.to_string()}"
                    
#                 elif file_name.lower().endswith('.docx'):
#                     from docx import Document
#                     doc = Document(uploaded_file)
#                     for paragraph in doc.paragraphs:
#                         file_content += paragraph.text + "\n"
                        
#                 elif file_name.lower().endswith('.xlsx'):
#                     import pandas as pd
#                     df = pd.read_excel(uploaded_file)
#                     file_content = f"Excel 파일 내용:\n{df.to_string()}"
                
#                 # 파일 내용을 세션에 저장
#                 st.session_state.uploaded_file_content = file_content
#                 st.session_state.uploaded_file_name = file_name
#                 st.success(f"✅ {file_name} 파일이 성공적으로 업로드되었습니다!")
                
#             except Exception as e:
#                 st.error(f"❌ 파일 처리 중 오류가 발생했습니다: {str(e)}")
        
#             # 파일 삭제 버튼
#             if st.button("🗑️ 파일 삭제", key="delete_file"):
#                 del st.session_state.uploaded_file_content
#                 del st.session_state.uploaded_file_name
#                 st.rerun()
#         else:
#             st.markdown("**📋 지원 파일 형식:** PDF, TXT, CSV, DOCX, XLSX")
    

#         # 최근 대화 확인 섹션
#         st.markdown("---")
#         st.header("💬 최근 대화 확인")
        
#         # 최근 대화 토글 버튼
#         if st.session_state.get('show_recent_messages', False):
#             if st.button("❌ 닫기", key="show_recent_toggle"):
#                 st.session_state.show_recent_messages = False
#                 st.rerun()
#         else:
#             if st.button("📋 최근 대화 보기", key="show_recent_toggle"):
#                 st.session_state.show_recent_messages = True
#                 st.rerun()

#     # 메시지 표시
#     for sender, msg in st.session_state.messages:
#         UI.bubble(sender, msg)
 
#     # ---------- 최근 대화 표시 (큰 창) ----------
#     if st.session_state.get('show_recent_messages', False):
#         st.markdown("---")
#         st.markdown("## 📋 최근 대화 내역")
        
#         recent_messages = DB.get_recent_messages(limit=20)
        
#         if recent_messages:
#             for i, m in enumerate(reversed(recent_messages)):
#                 st.markdown("---")
#                 col1, col2 = st.columns([1, 4])
                
#                 with col1:
#                     st.markdown(f"**{m.created_at.strftime('%m-%d %H:%M')}**")
#                 with col2:
#                     if st.button(f"🔄 다시 질문하기", key=f"reask_recent_main_{i}"):
#                         st.session_state.messages.append(("user", m.user_text))
#                         st.session_state.messages.append(("bot", m.bot_text))
#                         st.rerun()
                
#                 with st.expander(f"💬 {m.user_text[:50]}{'...' if len(m.user_text) > 50 else ''}", expanded=False):
#                     st.markdown("**🙋 질문:**")
#                     st.markdown(f"{m.user_text}")
#                     st.markdown("**🤖 답변:**")
#                     st.markdown(f"{m.bot_text}")
                    
#                     if len(m.bot_text) > 200:
#                         st.markdown(f"*답변 길이: {len(m.bot_text)}자*")
             
#             st.markdown("---")
#             st.markdown(f"**총 {len(recent_messages)}개의 대화 내역이 있습니다.**")
            
#             col1, col2 = st.columns([1, 5])
#             with col1:
#                 if st.button("❌ 최근 대화 닫기", key="close_recent_main"):
#                     st.session_state.show_recent_messages = False
#                     st.rerun()
            
#             st.info("💡 위 내용을 스크롤하여 모든 대화 내역을 확인할 수 있습니다.")
#         else:
#             st.info("아직 대화 내역이 없습니다. 첫 번째 질문을 해보세요!")
            
#             col1, col2 = st.columns([1, 5])
#             with col1:
#                 if st.button("❌ 최근 대화 닫기", key="close_recent_main"):
#                     st.session_state.show_recent_messages = False
#                     st.rerun()
         
#         st.markdown("---")
 
#     # ---------- 자주 묻는 질문 (질문 입력창 바로 위에 고정) ----------
#     popular_questions = DB.get_user_popular_questions(limit=5)
#     st.markdown("#### 🔥 자주 묻는 질문")
    
#     if popular_questions:
#         cols = st.columns(5, gap="small")
#         for i, (question, answer, count) in enumerate(popular_questions):
#             with cols[i]:
#                 display_question = question[:20] + "..." if len(question) > 20 else question
#                 if st.button(
#                     f"{display_question}", 
#                     key=f"popular_btn_{i}", 
#                     use_container_width=True,
#                     help=f"질문: {question}\n답변: {answer[:100]}..."
#                 ):
#                     st.session_state.messages.append(("user", question))
#                     st.session_state.messages.append(("bot", answer))
#                     st.rerun()
#     else:
#         cols = st.columns(5, gap="small")
#         for i, faq in enumerate(BEXCO_FAQ[:5]):
#             with cols[i]:
#                 display_question = faq["question"][:20] + "..." if len(faq["question"]) > 20 else faq["question"]
#                 if st.button(
#                     f"{display_question}", 
#                     key=f"default_faq_{i}", 
#                     use_container_width=True,
#                     help=f"질문: {faq['question']}\n답변: {faq['answer'][:100]}..."
#                 ):
#                     st.session_state.messages.append(("user", faq["question"]))
#                     st.session_state.messages.append(("bot", faq["answer"]))
#                     st.rerun()
     
#     # ---------- 입력 ----------
#     if prompt := st.chat_input("벡스코에 대해 궁금한 점을 물어보세요..."):
#         st.session_state.messages.append(("user", prompt))
        
#         with st.spinner("💭 답변을 생성하고 있습니다..."):
#             reply = bot.reply(prompt)
        
#         st.session_state.messages.append(("bot", reply))
#         DB.save_qa(st.session_state.session_id, prompt, reply)
#         st.rerun()


# if __name__ == "__main__":
#     main()


import os
import asyncio
import streamlit as st
import google.generativeai as genai
from datetime import datetime

from config import Config
from utils import DB, UI
from bexco_data import BEXCO_FAQ, search_bexco_info
from add_partition import add_monthly_partition


# ======================= 벡스코 챗봇 =======================
class BexcoChatbot:
    @staticmethod
    @st.cache_resource
    def load_model():
        try:
            with open(Config.API_KEY_PATH, "r", encoding="utf-8") as f:
                api_key = f.read().strip()
            genai.configure(api_key=api_key)
            return genai.GenerativeModel(Config.MODEL_NAME)
        except Exception as e:
            st.error(f"모델 로드 실패: {e}")
            st.stop()

    def __init__(self):
        self.model = BexcoChatbot.load_model()
        self._ensure_session()

    @staticmethod
    def _ensure_session():
        DB.init()
        ss = st.session_state
        ss.setdefault("session_id", os.urandom(8).hex())
        ss.setdefault("chat_session", BexcoChatbot.load_model().start_chat(history=[]))
        ss.setdefault("messages", [])
        ss.setdefault("show_db_recent", False)

    def _create_bexco_context(self, user_prompt: str) -> str:
        context = f"사용자 질문: {user_prompt}\n\n간단하고 직접적으로 답변해주세요."
        if hasattr(st.session_state, 'uploaded_file_content') and st.session_state.uploaded_file_content:
            context += f"""

파일 내용:
{st.session_state.uploaded_file_content[:2000]}{'...' if len(st.session_state.uploaded_file_content) > 2000 else ''}
"""
        return context

    # ================= 병렬 처리 =================
    async def _dataset_lookup(self, query: str):
        return search_bexco_info(query)

    async def _model_lookup(self, query: str):
        context = self._create_bexco_context(query)
        stream = st.session_state.chat_session.send_message(context, stream=True)
        text = ""
        for chunk in stream:
            if hasattr(chunk, "text"):
                text += chunk.text
        return text

    async def reply_async(self, user_prompt: str) -> str:
        try:
            dataset_task = asyncio.create_task(self._dataset_lookup(user_prompt))
            model_task = asyncio.create_task(self._model_lookup(user_prompt))

            dataset_result = await dataset_task
            if dataset_result:  # ✅ 데이터셋에서 답변 있으면 즉시 반환
                model_task.cancel()
                return "\n".join(dataset_result[:3])
            else:               # ✅ 없으면 모델 결과 반환
                return await model_task
        except Exception as e:
            return f"응답 생성 중 오류가 발생했습니다: {e}"

    def reply(self, user_prompt: str) -> str:
        return asyncio.run(self.reply_async(user_prompt))


# ======================= 메인 =======================
def main():
    st.set_page_config(
        page_title="부산 벡스코 챗봇",
        layout="wide",
        initial_sidebar_state="expanded"
    )

    # ✅ 실행 시 매달 파티션 자동 생성
    try:
        add_monthly_partition()
    except Exception as e:
        st.warning(f"⚠️ 파티션 생성 중 오류 발생: {e}")

    # 헤더
    st.title("🏢 부산벡스코(BEXCO) 챗봇")

    # CSS 스타일 적용
    UI.css()

    # 챗봇 초기화
    bot = BexcoChatbot()

    # ---------- 사이드바 ----------
    with st.sidebar:
        st.header("📎 파일 업로드")
        st.markdown("파일을 업로드하고 내용에 대해 질문하세요!")

        uploaded_file = st.file_uploader(
            "파일 선택",
            type=["pdf", "txt", "csv", "docx", "xlsx"],
            help="PDF, TXT, CSV, DOCX, XLSX 파일 업로드 가능"
        )

        if uploaded_file is not None:
            file_content = ""
            file_name = uploaded_file.name
            try:
                if file_name.lower().endswith('.pdf'):
                    import PyPDF2
                    pdf_reader = PyPDF2.PdfReader(uploaded_file)
                    for page in pdf_reader.pages:
                        file_content += page.extract_text() + "\n"
                elif file_name.lower().endswith('.txt'):
                    file_content = uploaded_file.read().decode('utf-8')
                elif file_name.lower().endswith('.csv'):
                    import pandas as pd
                    df = pd.read_csv(uploaded_file)
                    file_content = f"CSV 파일 내용:\n{df.to_string()}"
                elif file_name.lower().endswith('.docx'):
                    from docx import Document
                    doc = Document(uploaded_file)
                    for paragraph in doc.paragraphs:
                        file_content += paragraph.text + "\n"
                elif file_name.lower().endswith('.xlsx'):
                    import pandas as pd
                    df = pd.read_excel(uploaded_file)
                    file_content = f"Excel 파일 내용:\n{df.to_string()}"

                st.session_state.uploaded_file_content = file_content
                st.session_state.uploaded_file_name = file_name
                st.success(f"✅ {file_name} 업로드 완료!")
            except Exception as e:
                st.error(f"❌ 파일 처리 오류: {str(e)}")

            if st.button("🗑️ 파일 삭제", key="delete_file"):
                del st.session_state.uploaded_file_content
                del st.session_state.uploaded_file_name
                st.rerun()
        else:
            st.markdown("**📋 지원 형식:** PDF, TXT, CSV, DOCX, XLSX")

        # 최근 대화 확인
        st.markdown("---")
        st.header("💬 최근 대화 확인")

        if st.session_state.get('show_recent_messages', False):
            if st.button("❌ 닫기", key="show_recent_toggle"):
                st.session_state.show_recent_messages = False
                st.rerun()
        else:
            if st.button("📋 최근 대화 보기", key="show_recent_toggle"):
                st.session_state.show_recent_messages = True
                st.rerun()

    # 메시지 표시
    for sender, msg in st.session_state.messages:
        UI.bubble(sender, msg)

    # 최근 대화 표시
    if st.session_state.get('show_recent_messages', False):
        st.markdown("---")
        st.markdown("## 📋 최근 대화 내역")
        recent_messages = DB.get_recent_messages(limit=20)

        if recent_messages:
            for i, m in enumerate(reversed(recent_messages)):
                st.markdown("---")
                col1, col2 = st.columns([1, 4])
                with col1:
                    st.markdown(f"**{m.created_at.strftime('%m-%d %H:%M')}**")
                with col2:
                    if st.button("🔄 다시 질문하기", key=f"reask_recent_main_{i}"):
                        st.session_state.messages.append(("user", m.user_text))
                        st.session_state.messages.append(("bot", m.bot_text))
                        st.rerun()
                with st.expander(f"💬 {m.user_text[:50]}{'...' if len(m.user_text) > 50 else ''}", expanded=False):
                    st.markdown("**🙋 질문:**")
                    st.markdown(f"{m.user_text}")
                    st.markdown("**🤖 답변:**")
                    st.markdown(f"{m.bot_text}")
        else:
            st.info("아직 대화 내역이 없습니다.")

    # ---------- 자주 묻는 질문 ----------
    popular_questions = DB.get_user_popular_questions(limit=5)
    st.markdown("#### 🔥 자주 묻는 질문")

    if popular_questions:
        cols = st.columns(5, gap="small")
        for i, (question, answer, count) in enumerate(popular_questions):
            with cols[i]:
                display_question = question[:20] + "..." if len(question) > 20 else question
                if st.button(f"{display_question}", key=f"popular_btn_{i}", use_container_width=True):
                    st.session_state.messages.append(("user", question))
                    st.session_state.messages.append(("bot", answer))
                    st.rerun()
    else:
        cols = st.columns(5, gap="small")
        for i, faq in enumerate(BEXCO_FAQ[:5]):
            with cols[i]:
                display_question = faq["question"][:20] + "..." if len(faq["question"]) > 20 else faq["question"]
                if st.button(f"{display_question}", key=f"default_faq_{i}", use_container_width=True):
                    st.session_state.messages.append(("user", faq["question"]))
                    st.session_state.messages.append(("bot", faq["answer"]))
                    st.rerun()

    # ---------- 입력 ----------
    if prompt := st.chat_input("벡스코에 대해 궁금한 점을 물어보세요..."):
        st.session_state.messages.append(("user", prompt))
        with st.spinner("💭 답변 생성 중..."):
            reply = bot.reply(prompt)
        st.session_state.messages.append(("bot", reply))
        DB.save_qa(st.session_state.session_id, prompt, reply)
        st.rerun()


if __name__ == "__main__":
    main()
